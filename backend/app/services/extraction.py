"""
Extraction Service — Phase 1
Handles born-digital PDFs (selectable text + tables) and DOCX files.
Scanned PDFs (OCR) are Phase 2.
"""
import logging
import tempfile
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pdfplumber
import fitz  # PyMuPDF — for image detection only in Phase 1
from docx import Document as DocxDocument
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


# ── Data Structures ───────────────────────────────────────────────────────────

@dataclass
class PageContent:
    page_number: int
    text_blocks: list[str]       = field(default_factory=list)  # pdfplumber text
    table_blocks: list[str]      = field(default_factory=list)  # markdown tables
    has_images: bool             = False  # detected but not processed in Phase 1
    source_type: str             = "text"  # 'text' | 'table' | 'mixed'


@dataclass
class ExtractedDocument:
    file_id: str
    filename: str
    file_type: str               # 'pdf' | 'docx'
    pages: list[PageContent]     = field(default_factory=list)


# ── Table Converter ───────────────────────────────────────────────────────────

def _table_to_markdown(table: list[list]) -> str:
    """Convert pdfplumber table rows to a Markdown table string."""
    if not table or not table[0]:
        return ""
    # Clean None values
    cleaned = [[str(cell).strip() if cell is not None else "" for cell in row] for row in table]
    header = cleaned[0]
    separator = ["---"] * len(header)
    rows = cleaned[1:] if len(cleaned) > 1 else []

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in rows:
        # Pad short rows
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(padded) + " |")
    return "\n".join(lines)


def _docx_table_to_markdown(table) -> str:
    """Convert a python-docx Table to a Markdown table string."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    header = rows[0]
    separator = ["---"] * len(header)
    body = rows[1:]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in body:
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(padded) + " |")
    return "\n".join(lines)


# ── PDF Extraction ────────────────────────────────────────────────────────────

def extract_pdf(file_path: str, file_id: str, filename: str) -> ExtractedDocument:
    """
    Extract content from a born-digital PDF.
    - pdfplumber: selectable text + tables
    - PyMuPDF: image detection (Phase 1 only — no captioning)
    - Scanned pages (no selectable text) are skipped with a warning.
    """
    doc = ExtractedDocument(file_id=file_id, filename=filename, file_type="pdf")
    skipped_pages = []

    # Open with PyMuPDF to detect images
    fitz_doc = fitz.open(file_path)

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_content = PageContent(page_number=page_num)

            # ── Text extraction ──────────────────────────────────────────────
            raw_text = page.extract_text()
            if raw_text and raw_text.strip():
                # Clean and split into logical blocks by double-newline
                blocks = [b.strip() for b in raw_text.split("\n\n") if b.strip()]
                page_content.text_blocks = blocks
            else:
                skipped_pages.append(page_num)
                logger.warning(
                    f"[{filename}] Page {page_num}: no selectable text found. "
                    "Scanned pages are not supported in Phase 1 (OCR is Phase 2). Skipping."
                )

            # ── Table extraction ─────────────────────────────────────────────
            tables = page.extract_tables() or []
            for table in tables:
                md = _table_to_markdown(table)
                if md:
                    page_content.table_blocks.append(md)

            # ── Image detection (Phase 1: detect only, no captioning) ────────
            fitz_page = fitz_doc[page_num - 1]
            image_list = fitz_page.get_images(full=True)
            if image_list:
                page_content.has_images = True

            # ── Determine source type ────────────────────────────────────────
            has_text   = bool(page_content.text_blocks)
            has_tables = bool(page_content.table_blocks)
            if has_text and has_tables:
                page_content.source_type = "mixed"
            elif has_tables:
                page_content.source_type = "table"
            else:
                page_content.source_type = "text"

            # Only append page if it has content
            if page_content.text_blocks or page_content.table_blocks:
                doc.pages.append(page_content)

    fitz_doc.close()

    if skipped_pages:
        logger.info(f"[{filename}] Skipped {len(skipped_pages)} page(s) with no selectable text: {skipped_pages}")

    logger.info(f"[{filename}] PDF extraction complete. {len(doc.pages)} page(s) with content.")
    return doc


# ── DOCX Extraction ───────────────────────────────────────────────────────────

def extract_docx(file_path: str, file_id: str, filename: str) -> ExtractedDocument:
    """
    Extract content from a Word (.docx) file using python-docx.
    Preserves headings, paragraphs, and tables.
    """
    doc = ExtractedDocument(file_id=file_id, filename=filename, file_type="docx")
    docx = DocxDocument(file_path)

    # DOCX has no pages — we treat the whole doc as page 1
    # and split into logical sections by heading boundaries
    page_content = PageContent(page_number=1, source_type="text")
    current_block_lines: list[str] = []

    def flush_block():
        nonlocal current_block_lines
        text = "\n".join(current_block_lines).strip()
        if text:
            page_content.text_blocks.append(text)
        current_block_lines = []

    for element in docx.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            para_text = "".join(node.text or "" for node in element.iter() if node.tag.endswith("t")).strip()
            if not para_text:
                flush_block()
                continue
            # Detect heading style
            style_el = element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle")
            if style_el is not None:
                style_val = style_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "")
                if "Heading" in style_val or "heading" in style_val:
                    flush_block()                   # start new block at each heading
                    current_block_lines.append(f"## {para_text}")
                    continue
            current_block_lines.append(para_text)

        elif tag == "tbl":
            flush_block()
            # Process table
            table_rows = []
            for row_el in element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"):
                row = []
                for cell_el in row_el.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc"):
                    cell_text = "".join(n.text or "" for n in cell_el.iter() if n.tag.endswith("t")).strip()
                    row.append(cell_text)
                if any(row):
                    table_rows.append(row)
            if table_rows:
                # Manual markdown conversion since we have raw element data
                if len(table_rows) >= 1:
                    header = table_rows[0]
                    sep    = ["---"] * len(header)
                    body   = table_rows[1:]
                    lines  = ["| " + " | ".join(header) + " |", "| " + " | ".join(sep) + " |"]
                    for row in body:
                        padded = row + [""] * (len(header) - len(row))
                        lines.append("| " + " | ".join(padded) + " |")
                    md_table = "\n".join(lines)
                    page_content.table_blocks.append(md_table)

    flush_block()

    if page_content.text_blocks or page_content.table_blocks:
        if page_content.text_blocks and page_content.table_blocks:
            page_content.source_type = "mixed"
        elif page_content.table_blocks:
            page_content.source_type = "table"
        doc.pages.append(page_content)

    logger.info(f"[{filename}] DOCX extraction complete. {len(page_content.text_blocks)} text blocks, {len(page_content.table_blocks)} tables.")
    return doc


# ── Main Entry Point ──────────────────────────────────────────────────────────

def extract_document(local_file_path: str, file_id: str, filename: str, file_type: str) -> ExtractedDocument:
    """
    Route to the correct extractor based on file type.
    Raises ValueError for unsupported types.
    """
    file_type = file_type.lower()
    if file_type == "pdf":
        return extract_pdf(local_file_path, file_id, filename)
    elif file_type == "docx":
        return extract_docx(local_file_path, file_id, filename)
    else:
        raise ValueError(f"Unsupported file type: '{file_type}'. Only 'pdf' and 'docx' are supported.")
